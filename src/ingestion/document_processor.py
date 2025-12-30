"""
Document processing for multiple file types
Handles text, PDF, DOCX, and images with OCR
"""

import os
from pathlib import Path
from typing import List, Dict, Any
import hashlib
from dataclasses import dataclass, field
from datetime import datetime

# Text processing
import PyPDF2
from docx import Document as DocxDocument
import pytesseract
from PIL import Image

# Logging
from src.utils.logger import setup_logger

logger = setup_logger(__name__)


@dataclass
class Document:
    """Represents a processed document"""
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    doc_id: str = field(default="")
    
    def __post_init__(self):
        """Generate doc_id if not provided"""
        if not self.doc_id:
            # Generate hash from content and source
            source = self.metadata.get('source', '')
            hash_input = f"{source}:{self.content[:1000]}"
            self.doc_id = hashlib.sha256(hash_input.encode()).hexdigest()[:16]


class DocumentProcessor:
    """Process documents of various types"""
    
    def __init__(self, config):
        """Initialize document processor"""
        self.config = config
        self.supported_extensions = {
            '.txt': self._process_text,
            '.md': self._process_text,
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.png': self._process_image,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
        }
    
    def process_file(self, file_path: str) -> List[Document]:
        """
        Process a single file and return documents
        
        Args:
            file_path: Path to the file
            
        Returns:
            List of Document objects
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension not in self.supported_extensions:
            logger.warning(f"Unsupported file type: {extension}")
            return []
        
        try:
            processor = self.supported_extensions[extension]
            documents = processor(path)
            
            # Add common metadata
            for doc in documents:
                doc.metadata.update({
                    'source': str(path),
                    'filename': path.name,
                    'file_type': extension[1:],
                    'processed_at': datetime.now().isoformat(),
                    'file_size': path.stat().st_size
                })
            
            return documents
            
        except Exception as e:
            logger.error(f"Error processing {file_path}: {str(e)}")
            raise
    
    def _process_text(self, path: Path) -> List[Document]:
        """Process text/markdown files"""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Clean content
            content = self._clean_text(content)
            
            if not content.strip():
                logger.warning(f"Empty content in {path}")
                return []
            
            return [Document(content=content)]
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(path, 'r', encoding=encoding) as f:
                        content = f.read()
                    return [Document(content=self._clean_text(content))]
                except:
                    continue
            
            raise ValueError(f"Could not decode {path} with any encoding")
    
    def _process_pdf(self, path: Path) -> List[Document]:
        """Process PDF files"""
        documents = []
        
        try:
            with open(path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        
                        if text and text.strip():
                            text = self._clean_text(text)
                            
                            doc = Document(
                                content=text,
                                metadata={'page': page_num + 1}
                            )
                            documents.append(doc)
                    
                    except Exception as e:
                        logger.error(f"Error extracting page {page_num + 1} from {path}: {str(e)}")
            
            if not documents:
                logger.warning(f"No text extracted from PDF: {path}")
            
            return documents
            
        except Exception as e:
            logger.error(f"Error processing PDF {path}: {str(e)}")
            raise
    
    def _process_docx(self, path: Path) -> List[Document]:
        """Process DOCX files"""
        try:
            doc = DocxDocument(path)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        paragraphs.append(row_text)
            
            content = '\n\n'.join(paragraphs)
            content = self._clean_text(content)
            
            if not content.strip():
                logger.warning(f"Empty content in {path}")
                return []
            
            return [Document(content=content)]
            
        except Exception as e:
            logger.error(f"Error processing DOCX {path}: {str(e)}")
            raise
    
    def _process_image(self, path: Path) -> List[Document]:
        """Process image files with OCR"""
        if not self.config.ocr_enabled:
            logger.info(f"OCR disabled, skipping image: {path}")
            return []
        
        try:
            # Open and process image
            image = Image.open(path)
            
            # Convert to RGB if necessary
            if image.mode not in ('RGB', 'L'):
                image = image.convert('RGB')
            
            # Perform OCR
            custom_config = r'--oem 3 --psm 6'
            text = pytesseract.image_to_string(
                image,
                lang='+'.join(self.config.ocr_languages),
                config=custom_config
            )
            
            text = self._clean_text(text)
            
            if not text.strip():
                logger.warning(f"No text extracted from image: {path}")
                return []
            
            # Also extract basic image info
            width, height = image.size
            metadata = {
                'image_width': width,
                'image_height': height,
                'image_mode': image.mode,
                'ocr_extracted': True
            }
            
            return [Document(content=text, metadata=metadata)]
            
        except Exception as e:
            logger.error(f"Error processing image {path}: {str(e)}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n')]
        lines = [line for line in lines if line]
        
        # Join with single newline
        text = '\n'.join(lines)
        
        # Remove excessive newlines
        while '\n\n\n' in text:
            text = text.replace('\n\n\n', '\n\n')
        
        return text.strip()
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions"""
        return list(self.supported_extensions.keys())
